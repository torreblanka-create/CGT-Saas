"""
revision_direccion.py — Registro y Acta de Revisión por la Dirección (ISO 9001, 14001, 45001).
Permite documentar las entradas, salidas y generar un acta automática en formato Word (.docx).
"""
import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except (ImportError, ModuleNotFoundError):
    DOCX_AVAILABLE = False
    Document = None

from src.infrastructure.database import ejecutar_query, obtener_dataframe
from core.utils import is_valid_context, show_context_warning, render_hybrid_date_input

def _generar_acta_word(datos):
    """Genera un archivo Word usando python-docx en memoria."""
    doc = Document()
    
    # Estilos del título
    styles = doc.styles
    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0, 83, 156)
    
    # Encabezado
    titulo = doc.add_heading(f"Acta de Revisión por la Dirección SGI", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Metadatos de la reunión
    p_meta = doc.add_paragraph()
    p_meta.add_run("Fecha de Reunión: ").bold = True
    p_meta.add_run(f"{datos['fecha']}\n")
    p_meta.add_run("Empresa / Contrato: ").bold = True
    p_meta.add_run(f"{datos['empresa']} / {datos['contrato']}\n")
    p_meta.add_run("Periodo Evaluado: ").bold = True
    p_meta.add_run(f"{datos['periodo']}\n")
    
    # Participantes
    doc.add_heading("1. Participantes", level=2)
    doc.add_paragraph(datos['participantes'])
    
    # Entradas de la revisión (Inputs ISO)
    doc.add_heading("2. Entradas de la Revisión (Análisis de Desempeño)", level=2)
    doc.add_paragraph("a) Estado de las acciones de revisiones anteriores: " + datos['estado_acciones'])
    doc.add_paragraph("b) Cambios en cuestiones externas/internas y partes interesadas: " + datos['cambios_contexto'])
    doc.add_paragraph("c) Desempeño y eficacia del SGI (Auditorías, NCRs, KPIs): " + datos['desempeno_sgi'])
    doc.add_paragraph("d) Adecuación de los recursos: " + datos['adecuacion_recursos'])
    doc.add_paragraph("e) Evaluación de riesgos y oportunidades: " + datos['riesgos_ops'])
    doc.add_paragraph("f) Oportunidades de Mejora Continua: " + datos['oportunidades'])
    
    # Salidas de la revisión (Outputs ISO)
    doc.add_heading("3. Salidas de la Revisión (Decisiones y Acuerdos)", level=2)
    doc.add_paragraph(datos['acuerdos'])
    
    # Firmas
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    
    p1 = row[0].paragraphs[0]
    p1.add_run("______________________________\n").bold = True
    p1.add_run("Firma Alta Dirección")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p2 = row[1].paragraphs[0]
    p2.add_run("______________________________\n").bold = True
    p2.add_run("Firma Responsable SGI")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Guardar a buffer en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def render_revision_direccion(DB_PATH, filtros):
    st.markdown("<h2 style='color: var(--cgt-blue);'>🤝 Revisiones por la Dirección (Actas ISO)</h2>", unsafe_allow_html=True)
    st.write("Cumplimiento del requisito normativo (cláusula 9.3) para evaluar la conveniencia, adecuación y eficacia del Sistema de Gestión Integrado.")

    if not DOCX_AVAILABLE:
        st.warning("⚠️ Esta función requiere la librería 'python-docx' que no está disponible.")
        return

    if not is_valid_context(filtros):
        show_context_warning()
        return

    emp_id = filtros.get('empresa_id')
    con_id = filtros.get('contrato_id')
    empresa_nom = filtros.get('empresa_nom', 'N/A')
    contrato_nom = filtros.get('contrato_nom', 'N/A')

    tab_reg, tab_hist = st.tabs(["\U0001f4dd Registrar Acta", "\U0001f5c3\ufe0f Historial y Exportaci\u00f3n"])

    with tab_reg:
        st.markdown("#### Formulario Estructurado de Entradas y Salidas SGI")

        with st.form("form_revision_dir"):
            col1, col2 = st.columns(2)
            with col1:
                f_rev = render_hybrid_date_input("Fecha de Reuni\u00f3n", key="f_rev_dir", value=datetime.now().date())
            with col2:
                periodo = st.text_input("Periodo Analizado", placeholder="Ej: A\u00f1o 2026", key="per_rev")

            st.markdown("#### 1. Informaci\u00f3n General")
            part = st.text_area("Participantes (Nombre y Cargo)", height=80, placeholder="Juan P\u00e9rez - Gerente General\nMar\u00eda L\u00f3pez - Jefe SGI")

            st.markdown("#### 2. Entradas de la Revisi\u00f3n")
            col_a, col_b = st.columns(2)
            with col_a:
                inp1 = st.text_area("a) Estado de las acciones de revisiones anteriores", height=100)
                inp2 = st.text_area("b) Cambios en el contexto (interno/externo) y partes interesadas", height=100)
                inp3 = st.text_area("c) Desempe\u00f1o del SGI (Grado cumplimiento Objetivos, Auditor\u00edas, Incidentes, Quejas Cliente, NCR)", height=100)
            with col_b:
                inp4 = st.text_area("d) Adecuaci\u00f3n de Recursos (Humanos, Infraestructura, Tecnol\u00f3gicos)", height=100)
                inp5 = st.text_area("e) Desempe\u00f1o de proveedores y riesgos/oportunidades", height=100)
                inp6 = st.text_area("f) Oportunidades de Mejora Continua", height=100)

            st.markdown("#### 3. Salidas de la Revisi\u00f3n (Decisiones)")
            acuerdos = st.text_area("Acuerdos, Decisiones y Asignaci\u00f3n de Recursos para la Mejora (Plan de Acci\u00f3n)", height=150)

            if st.form_submit_button("\u2705 Guardar y Generar Registro", type="primary", use_container_width=True):
                import json
                datos_json = json.dumps({
                    "periodo": periodo,
                    "estado_acciones": inp1,
                    "cambios_contexto": inp2,
                    "desempeno_sgi": inp3,
                    "adecuacion_recursos": inp4,
                    "riesgos_ops": inp5,
                    "oportunidades": inp6
                })
                query = """
                    INSERT INTO sgi_revision_direccion
                    (fecha, participantes, acuerdos, estado, empresa_id, contrato_id)
                    VALUES (?, ?, ?, ?, ?, ?)
                """
                ejecutar_query(DB_PATH, query, (str(f_rev), part, acuerdos, datos_json, emp_id, con_id), commit=True)
                st.success("\u2705 Acta de Revisi\u00f3n guardada correctamente. Puede descargar el Word en el Historial.")
                st.balloons()
                st.rerun()

    with tab_hist:
        st.markdown("#### Historial de Actas Emitidas")
        df_rev = obtener_dataframe(DB_PATH, "SELECT id, fecha, participantes, acuerdos, estado FROM sgi_revision_direccion WHERE empresa_id=? ORDER BY id DESC", (emp_id,))
        
        if df_rev.empty:
            st.info("Aún no hay actas registradas para esta empresa.")
        else:
            # Mostrar tabla resumen
            st.dataframe(df_rev[['id', 'fecha', 'participantes']], use_container_width=True, hide_index=True)
            
            st.markdown("#### 📄 Descargar Acta")
            c_sel, c_btn, _ = st.columns([2, 2, 2])
            
            id_sel = c_sel.selectbox("Seleccione ID de Acta:", df_rev['id'].tolist())
            
            if id_sel:
                row = df_rev[df_rev['id'] == id_sel].iloc[0]
                
                # Parsear el JSON oculto en "estado" para reconstruir el Word
                import json
                try:
                    metadata = json.loads(row['estado'])
                except:
                    metadata = {
                        "periodo": "No especificado",
                        "estado_acciones": "N/A", "cambios_contexto": "N/A", "desempeno_sgi": "N/A",
                        "adecuacion_recursos": "N/A", "riesgos_ops": "N/A", "oportunidades": "N/A"
                    }
                
                datos_docx = {
                    "fecha": row['fecha'],
                    "empresa": empresa_nom,
                    "contrato": contrato_nom,
                    "periodo": metadata.get('periodo', 'N/A'),
                    "participantes": row['participantes'],
                    "estado_acciones": metadata.get('estado_acciones', 'N/A'),
                    "cambios_contexto": metadata.get('cambios_contexto', 'N/A'),
                    "desempeno_sgi": metadata.get('desempeno_sgi', 'N/A'),
                    "adecuacion_recursos": metadata.get('adecuacion_recursos', 'N/A'),
                    "riesgos_ops": metadata.get('riesgos_ops', 'N/A'),
                    "oportunidades": metadata.get('oportunidades', 'N/A'),
                    "acuerdos": row['acuerdos']
                }
                
                buffer_docx = _generar_acta_word(datos_docx)
                
                c_btn.download_button(
                    label="📥 Descargar Documento Word (.docx)",
                    data=buffer_docx,
                    file_name=f"Acta_Rev_Direccion_SGI_{id_sel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
                
                with st.expander("👁️ Ver acuerdos principales", expanded=True):
                    st.info(f"**Decisiones Tomadas:**\n\n{row['acuerdos']}")
